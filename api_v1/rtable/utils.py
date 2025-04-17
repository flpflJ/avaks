import os
import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

def createDocxSAU(raw_text: str, id):
    doc = Document()
    doc.add_heading('Опросник САУ', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.autofit = True

    for column in table.columns:
        column.width = Inches(3)

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Параметр, вопрос'
    hdr_cells[1].text = 'Ответ'
    hdr_cells[2].text = 'Примечания'


    data = []
    lines = raw_text.strip().split("\n")
    for line in lines:
        if line.strip():
            parts = line.split(":", 1)
            if len(parts) == 2:
                question_number = parts[0].strip().split(".")[0]
                question = parts[0].split(".", 1)[1].strip()
                answer = parts[1].strip()
                data.append((f"{question_number}. {question}", answer, ""))
    for item in data:
        row = table.add_row().cells
        row[0].text = item[0]
        row[1].text = item[1]
        row[2].text = item[2]

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(10)
    os.makedirs("docs", exist_ok=True)
    doc.save(f"docs/Опросник_САУ_готовый_{id}.docx")

def parse_raw_data(raw_text):
    data = {}
    current_main = None
    lines = raw_text.strip().split('\n')

    main_pattern = re.compile(r'^(\d+)\.\s+(.*?):\s*(.*)')
    sub_pattern = re.compile(r'^\s{4}(\d+)\.\s+(.*?):\s*(.*)')
    unit_pattern = re.compile(r'(.+?)\s*([кмчгКМГбБЦцA-Za-z/°²"%-]+)$')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        main_match = main_pattern.match(line)
        if main_match:
            num, question, answer = main_match.groups()
            current_main = num

            unit = '-'
            if any(x in question for x in ['не менее', 'не более']):
                unit_match = unit_pattern.search(answer)
                if unit_match:
                    answer, unit = unit_match.groups()

            data[num] = {
                'question': question,
                'unit': unit.strip(),
                'answer': answer.replace(unit, '').strip(),
                'subitems': []
            }
            continue

        sub_match = sub_pattern.match(line)
        if sub_match and current_main:
            _, sub_num, sub_question, sub_answer = sub_match.groups()
            data[current_main]['subitems'].append(
                (sub_num, sub_question, sub_answer)
            )

    return data


def create_survey_document(data, id):
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    widths = [Inches(0.6), Inches(3.5), Inches(1), Inches(2)]
    for i, width in enumerate(widths):
        table.columns[i].width = width

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№п/п'
    hdr_cells[1].text = 'Запрашиваемые данные'
    hdr_cells[2].text = 'Ед. изм.'
    hdr_cells[3].text = 'Технические характеристики'

    for num in range(1, 24):
        key = str(num)
        item = data.get(key, {'question': '', 'unit': '-', 'answer': '', 'subitems': []})

        row = table.add_row().cells
        row[0].text = key
        row[1].text = item['question']
        row[2].text = item['unit']
        row[3].text = item['answer']

        for sub in item['subitems']:
            sub_row = table.add_row().cells
            sub_row[0].text = ''
            sub_row[1].text = f"{sub[0]}. {sub[1]}"
            sub_row[2].text = '-'
            sub_row[3].text = sub[2]

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.style.font.name = 'Arial'
                paragraph.style.font.size = Pt(10)
                if cell.text.startswith(('1.', '2.', '3.', '4.', '5.')):
                    paragraph.paragraph_format.left_indent = Inches(0.3)

    os.makedirs("docs", exist_ok=True)
    doc.save(f"docs/Опросный_лист_заказчика_{id}.docx")

def split_text(text):
    split_index = text.find("1. Наименование предприятия-заказчика:")
    part1 = text[:split_index].strip()
    part2 = text[split_index:].strip()
    return part1,part2